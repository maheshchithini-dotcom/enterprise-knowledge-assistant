import os
import uuid
import pypdfium2 as pdfium
from docx import Document
from langchain_core.documents import Document as LangchainDoc
from langchain_text_splitters import RecursiveCharacterTextSplitter
from dotenv import load_dotenv

load_dotenv()

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", 1000))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", 200))


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pypdfium2."""
    text = ""
    pdf = pdfium.PdfDocument(file_path)
    for page_num in range(len(pdf)):
        page = pdf[page_num]
        textpage = page.get_textpage()
        page_text = textpage.get_text_range()
        if page_text.strip():
            text += f"\n--- Page {page_num + 1} ---\n{page_text}"
    pdf.close()
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a Word document."""
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text += row_text + "\n"
    return text.strip()


def process_document(file_path: str, filename: str):
    """
    Parse a PDF or DOCX file and split into LangChain Document chunks.
    Returns a tuple of (list of LangchainDoc, doc_id).
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        raw_text = extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        raw_text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    if not raw_text:
        raise ValueError(f"No text could be extracted from {filename}")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", " ", ""],
    )

    chunks = splitter.split_text(raw_text)

    doc_id = str(uuid.uuid4())
    documents = [
        LangchainDoc(
            page_content=chunk,
            metadata={
                "source": filename,
                "doc_id": doc_id,
                "chunk_index": i,
                "file_type": ext.lstrip("."),
            },
        )
        for i, chunk in enumerate(chunks)
    ]

    return documents, doc_id